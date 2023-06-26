'''
Create a program to accept the Dog breed and display related information that breed pull the data from the Dog
API and put it into Word file the file name should the the Dog breed name which is input from the user
'''
import requests
from docx import Document


def get_dog_breed_info(self):
    url = f"https://api.thedogapi.com/v1/breeds/search?q={breed}"
    response = requests.get(url)
    data = response.json()
    print(data)
    return data


breed = input("Enter the dog breed: ")
breed_info = get_dog_breed_info(breed)
document = Document()

if breed_info:
    name = breed_info[0]['name']
    weight = breed_info[0]['weight']
    height = breed_info[0]['height']
    bred_for = breed_info[0]['bred_for']
    breed_group = breed_info[0]['breed_group']
    life_span = breed_info[0]['life_span']
    temperament = breed_info[0]['temperament']
    reference_image_id = breed_info[0]['reference_image_id']

    document.add_heading(name, level=1)
    document.add_paragraph(f"Weight: {weight}")
    document.add_paragraph(f"Height: {height}")
    document.add_paragraph(f"Bred for: {bred_for}")
    document.add_paragraph(f"Breed group: {breed_group}")
    document.add_paragraph(f"Life_ span: {life_span}")
    document.add_paragraph(f"Temperament: {temperament}")
    document.add_paragraph(f"Reference Image Id: {reference_image_id}")

else:
    document.add_paragraph("Breed not found.")

document.save(f"{breed}_info.docx")
